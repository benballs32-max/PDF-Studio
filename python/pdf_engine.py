"""PDF processing engine — called as a subprocess by Electron main process."""
import sys
import json
import os
import tempfile

def _save(doc, src: str, dst: str, **save_kwargs):
    """Save fitz doc to dst. When dst == src (in-place), writes to a temp file
    then atomically replaces the original (required by PyMuPDF).
    Extra kwargs (e.g. encryption, permissions) are forwarded to doc.save()."""
    if os.path.normcase(os.path.abspath(src)) == os.path.normcase(os.path.abspath(dst)):
        fd, tmp = tempfile.mkstemp(suffix='.pdf', dir=os.path.dirname(os.path.abspath(dst)))
        try:
            os.close(fd)
            doc.save(tmp, **save_kwargs)
            doc.close()
            os.replace(tmp, dst)
        except Exception:
            try: os.unlink(tmp)
            except OSError: pass
            raise
    else:
        doc.save(dst, **save_kwargs)

def to_text(input: str, output: str, **_):
    import fitz
    doc = fitz.open(input)
    text = "\n\n".join(page.get_text() for page in doc)
    with open(output, "w", encoding="utf-8") as f:
        f.write(text)
    return {"success": True, "pages": len(doc)}

def to_html(input: str, output: str, **_):
    import fitz
    doc = fitz.open(input)
    parts = [page.get_text("html") for page in doc]
    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Converted PDF</title></head>
<body>{'<hr>'.join(parts)}</body></html>"""
    with open(output, "w", encoding="utf-8") as f:
        f.write(html)
    return {"success": True, "pages": len(doc)}

def to_image(input: str, output: str, format: str = "png", dpi: int = 150, **_):
    import fitz
    doc = fitz.open(input)
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    base, ext = os.path.splitext(output)
    saved = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat, alpha=False)
        path = f"{base}_p{i+1}{ext}" if len(doc) > 1 else output
        pix.save(path)
        saved.append(path)
    return {"success": True, "files": saved}

def to_docx(input: str, output: str, **_):
    import fitz
    from docx import Document
    from docx.shared import Pt

    doc = fitz.open(input)
    word = Document()
    for page in doc:
        for block in page.get_text("blocks"):
            text = block[4].strip()
            if text:
                para = word.add_paragraph(text)
                para.paragraph_format.space_after = Pt(6)
        word.add_page_break()
    word.save(output)
    return {"success": True, "pages": len(doc)}

def to_xlsx(input: str, output: str, **_):
    import fitz
    import openpyxl

    doc = fitz.open(input)
    wb = openpyxl.Workbook()
    for page_num, page in enumerate(doc):
        ws = wb.create_sheet(title=f"Page {page_num + 1}")
        tables = page.find_tables()
        if tables.tables:
            for table in tables.tables:
                for r, row in enumerate(table.extract()):
                    for c, cell in enumerate(row):
                        ws.cell(row=r + 1, column=c + 1, value=cell)
        else:
            for i, line in enumerate(page.get_text().splitlines()):
                ws.cell(row=i + 1, column=1, value=line)
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]
    wb.save(output)
    return {"success": True, "pages": len(doc)}

def merge(inputs: list, output: str, **_):
    import fitz
    result = fitz.open()
    for path in inputs:
        result.insert_pdf(fitz.open(path))
    result.save(output)
    return {"success": True, "pages": len(result)}

def split(input: str, output: str, **_):
    import fitz
    doc = fitz.open(input)
    os.makedirs(output, exist_ok=True)
    saved = []
    for i in range(len(doc)):
        out = fitz.open()
        out.insert_pdf(doc, from_page=i, to_page=i)
        path = os.path.join(output, f"page_{i+1}.pdf")
        out.save(path)
        saved.append(path)
    return {"success": True, "files": saved}

def compress(input: str, output: str, **_):
    import fitz
    doc = fitz.open(input)
    doc.save(output, garbage=4, deflate=True, clean=True)
    before = os.path.getsize(input)
    after = os.path.getsize(output)
    return {"success": True, "before_kb": before // 1024, "after_kb": after // 1024}

def rotate_page(input: str, output: str, page: int, angle: int, **_):
    import fitz
    doc = fitz.open(input)
    doc[page].set_rotation((doc[page].rotation + angle) % 360)
    _save(doc, input, output)
    return {"success": True}

def delete_page(input: str, output: str, page: int, **_):
    import fitz
    doc = fitz.open(input)
    doc.delete_page(page)
    count = len(doc)
    _save(doc, input, output)
    return {"success": True, "pages": count}

def apply_annotations(input: str, output: str, annotations: list, **_):
    import fitz

    def hex_to_rgb(h: str):
        h = h.lstrip('#')
        return tuple(int(h[i:i+2], 16) / 255 for i in (0, 2, 4))

    doc = fitz.open(input)
    for a in annotations:
        pg = doc[a['page']]
        color = hex_to_rgb(a.get('color', '#FACC15'))
        if a['type'] == 'highlight':
            annot = pg.add_highlight_annot(fitz.Rect(a['rect']))
            annot.set_colors(stroke=color); annot.update()
        elif a['type'] == 'text':
            annot = pg.add_text_annot(fitz.Point(a['x'], pg.rect.height - a['y']), a['content'])
            annot.set_colors(stroke=color); annot.update()
        elif a['type'] == 'ink':
            ink = [[fitz.Point(p[0], pg.rect.height - p[1]) for p in a['points']]]
            annot = pg.add_ink_annot(ink)
            annot.set_colors(stroke=color)
            annot.set_border(width=a.get('width', 2))
            annot.update()
    doc.save(output)
    return {"success": True}

def reorder_pages(input: str, output: str, order: list, **_):
    import fitz
    doc = fitz.open(input)
    doc.select(order)
    count = len(doc)
    _save(doc, input, output)
    return {"success": True, "pages": count}

def extract_pages(input: str, output: str, pages: list, **_):
    import fitz
    doc = fitz.open(input)
    out = fitz.open()
    for pg in sorted(set(pages)):
        out.insert_pdf(doc, from_page=pg, to_page=pg)
    out.save(output)
    return {"success": True, "pages": len(out)}

def insert_pages(input: str, output: str, source: str, position: int, **_):
    import fitz
    doc = fitz.open(input)
    src = fitz.open(source)
    doc.insert_pdf(src, start_at=position)
    src.close()
    count = len(doc)
    _save(doc, input, output)
    return {"success": True, "pages": count}

def crop_page(input: str, output: str, page: int, rect: list, **_):
    import fitz
    doc = fitz.open(input)
    doc[page].set_cropbox(fitz.Rect(rect))
    _save(doc, input, output)
    return {"success": True}

_STAMP_COLORS = {
    'APPROVED': (0.0, 0.55, 0.0),
    'DRAFT':    (0.75, 0.45, 0.0),
    'CONFIDENTIAL': (0.82, 0.0, 0.0),
    'COPY':     (0.0, 0.2, 0.8),
    'VOID':     (0.45, 0.0, 0.0),
}

def images_to_pdf(images: list, output: str, **_):
    import fitz
    doc = fitz.open()
    for img_path in images:
        imgdoc = fitz.open(img_path)
        pdfbytes = imgdoc.convert_to_pdf()
        imgpdf = fitz.open("pdf", pdfbytes)
        doc.insert_pdf(imgpdf)
        imgdoc.close(); imgpdf.close()
    doc.save(output)
    return {"success": True, "pages": len(doc)}

def office_to_pdf(input: str, output: str, **_):
    import shutil, subprocess, os
    for exe_name in ('libreoffice', 'soffice'):
        exe = shutil.which(exe_name)
        if exe:
            out_dir = os.path.dirname(os.path.abspath(output))
            subprocess.run([exe, '--headless', '--convert-to', 'pdf',
                            '--outdir', out_dir, input],
                           check=True, capture_output=True, timeout=120)
            base = os.path.splitext(os.path.basename(input))[0]
            generated = os.path.join(out_dir, base + '.pdf')
            if os.path.abspath(generated) != os.path.abspath(output):
                os.replace(generated, output)
            return {"success": True}
    try:
        from docx2pdf import convert as d2p
        d2p(input, output)
        return {"success": True}
    except ImportError:
        pass
    raise RuntimeError(
        "No Office conversion tool found. "
        "Install LibreOffice (free) or run: pip install docx2pdf (requires Microsoft Office)"
    )

def url_to_pdf(url: str, output: str, **_):
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(url, wait_until='networkidle', timeout=30000)
            page.pdf(path=output, format='A4', print_background=True)
            browser.close()
        return {"success": True}
    except ImportError:
        pass
    try:
        from weasyprint import HTML
        HTML(url).write_pdf(output)
        return {"success": True}
    except ImportError:
        pass
    raise RuntimeError(
        "No web-to-PDF tool found. Install one of:\n"
        "  pip install playwright && playwright install chromium\n"
        "  pip install weasyprint"
    )

def add_watermark(input: str, output: str, text: str, opacity: float = 0.3,
                  angle: float = 45, font_size: float = 60, color: list = None, **_):
    import fitz
    col = tuple(color) if color else (0.65, 0.65, 0.65)
    doc = fitz.open(input)
    for page in doc:
        r = page.rect
        center = fitz.Point(r.width / 2, r.height / 2)
        # Offset start so text appears centred around the page centre after rotation
        half_w = len(text) * font_size * 0.28
        start = fitz.Point(center.x - half_w, center.y + font_size * 0.25)
        page.insert_text(
            start, text,
            fontname="helv", fontsize=font_size,
            color=col,
            morph=(center, fitz.Matrix(angle)),
            overlay=True,
        )
    _save(doc, input, output)
    return {"success": True}

def add_stamp(input: str, output: str, stamp: str, position: str = 'top-right',
              page_num: int = -1, **_):
    import fitz
    stamp = stamp.upper()
    color = _STAMP_COLORS.get(stamp, (0.3, 0.3, 0.3))
    font_size = 30
    pad = 7
    doc = fitz.open(input)
    pages = [doc[page_num]] if 0 <= page_num < len(doc) else list(doc)
    for page in pages:
        r = page.rect
        # Estimate text box size (hebo is bold; chars ~0.62× fontsize wide)
        bw = len(stamp) * font_size * 0.62 + pad * 2
        bh = font_size + pad * 2
        if position == 'top-right':
            x, y = r.width - bw - 20, 18
        elif position == 'top-left':
            x, y = 20, 18
        elif position == 'bottom-right':
            x, y = r.width - bw - 20, r.height - bh - 18
        elif position == 'bottom-left':
            x, y = 20, r.height - bh - 18
        else:  # center
            x, y = (r.width - bw) / 2, (r.height - bh) / 2
        box = fitz.Rect(x, y, x + bw, y + bh)
        page.draw_rect(box, color=color, width=2, overlay=True)
        page.insert_text(fitz.Point(x + pad, y + font_size + pad * 0.4),
                         stamp, fontname="hebo", fontsize=font_size,
                         color=color, overlay=True)
    _save(doc, input, output)
    return {"success": True}

def add_header_footer(input: str, output: str, header: str = '', footer: str = '',
                      font_size: float = 10, **_):
    import fitz
    doc = fitz.open(input)
    for page in doc:
        r = page.rect
        if header:
            hw = len(header) * font_size * 0.55
            page.insert_text(fitz.Point(r.width / 2 - hw / 2, font_size + 6),
                              header, fontname="helv", fontsize=font_size,
                              color=(0, 0, 0), overlay=True)
            page.draw_line(fitz.Point(36, font_size + 10),
                           fitz.Point(r.width - 36, font_size + 10),
                           color=(0.7, 0.7, 0.7), width=0.5, overlay=True)
        if footer:
            fw = len(footer) * font_size * 0.55
            page.draw_line(fitz.Point(36, r.height - font_size - 10),
                           fitz.Point(r.width - 36, r.height - font_size - 10),
                           color=(0.7, 0.7, 0.7), width=0.5, overlay=True)
            page.insert_text(fitz.Point(r.width / 2 - fw / 2, r.height - 6),
                              footer, fontname="helv", fontsize=font_size,
                              color=(0, 0, 0), overlay=True)
    _save(doc, input, output)
    return {"success": True}

def add_page_numbers(input: str, output: str, position: str = 'bottom-center',
                     start: int = 1, font_size: float = 11, **_):
    import fitz
    doc = fitz.open(input)
    for i, page in enumerate(doc):
        r = page.rect
        text = str(start + i)
        tw = len(text) * font_size * 0.6
        if position == 'bottom-center':
            pt = fitz.Point(r.width / 2 - tw / 2, r.height - 8)
        elif position == 'bottom-right':
            pt = fitz.Point(r.width - tw - 20, r.height - 8)
        elif position == 'bottom-left':
            pt = fitz.Point(20, r.height - 8)
        elif position == 'top-center':
            pt = fitz.Point(r.width / 2 - tw / 2, font_size + 6)
        elif position == 'top-right':
            pt = fitz.Point(r.width - tw - 20, font_size + 6)
        else:  # top-left
            pt = fitz.Point(20, font_size + 6)
        page.insert_text(pt, text, fontname="helv", fontsize=font_size, overlay=True)
    _save(doc, input, output)
    return {"success": True}

def get_outline(input: str, **_):
    import fitz
    doc = fitz.open(input)
    toc = doc.get_toc()  # [[level, title, page_1based], ...]
    return {"success": True, "outline": [{"level": t[0], "title": t[1], "page": t[2]} for t in toc]}

def get_form_fields(input: str, **_):
    import fitz
    doc = fitz.open(input)
    fields = []
    for page_num, page in enumerate(doc):
        for widget in page.widgets():
            fields.append({
                'name': widget.field_name,
                'type': widget.field_type_string,
                'value': str(widget.field_value or ''),
                'page': page_num,
                'rect': list(widget.rect),
                'choices': list(widget.choice_values or []),
            })
    return {"success": True, "fields": fields}

def fill_form_fields(input: str, output: str, fields: dict, **_):
    import fitz
    doc = fitz.open(input)
    for page in doc:
        for widget in page.widgets():
            if widget.field_name in fields:
                widget.field_value = fields[widget.field_name]
                widget.update()
    _save(doc, input, output)
    return {"success": True}

def add_form_field(input: str, output: str, page: int, field_type: str,
                   name: str, rect: list, value: str = '', choices: list = None, **_):
    import fitz
    type_map = {
        'text':     fitz.PDF_WIDGET_TYPE_TEXT,
        'checkbox': fitz.PDF_WIDGET_TYPE_CHECKBOX,
        'radio':    fitz.PDF_WIDGET_TYPE_RADIOBUTTON,
        'dropdown': fitz.PDF_WIDGET_TYPE_COMBOBOX,
        'listbox':  fitz.PDF_WIDGET_TYPE_LISTBOX,
    }
    doc = fitz.open(input)
    pg = doc[page]
    widget = fitz.Widget()
    widget.field_type = type_map.get(field_type, fitz.PDF_WIDGET_TYPE_TEXT)
    widget.field_name = name
    widget.rect = fitz.Rect(rect)
    if value:
        widget.field_value = value
    if choices:
        widget.choice_values = choices
    pg.add_widget(widget)
    _save(doc, input, output)
    return {"success": True}

def export_form_data(input: str, output: str, format: str = 'json', **_):
    import fitz, json
    doc = fitz.open(input)
    data = {}
    for page in doc:
        for widget in page.widgets():
            data[widget.field_name] = str(widget.field_value or '')
    if format == 'csv':
        import csv
        with open(output, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Field Name', 'Value'])
            for k, v in data.items():
                writer.writerow([k, v])
    else:
        with open(output, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    return {"success": True, "field_count": len(data)}

def import_form_data(input: str, output: str, source: str, **_):
    import fitz, json, os
    ext = os.path.splitext(source)[1].lower()
    if ext == '.csv':
        import csv
        data = {}
        with open(source, encoding='utf-8') as f:
            for row in csv.DictReader(f):
                data[row['Field Name']] = row['Value']
    else:
        with open(source, encoding='utf-8') as f:
            data = json.load(f)
    doc = fitz.open(input)
    for page in doc:
        for widget in page.widgets():
            if widget.field_name in data:
                widget.field_value = data[widget.field_name]
                widget.update()
    _save(doc, input, output)
    return {"success": True}

def ocr_pdf(input: str, output: str, lang: str = 'eng', dpi: int = 300, **_):
    """Render every page to an image, run Tesseract OCR, and write a searchable PDF
    with an invisible text layer so Ctrl+F works on scanned documents."""
    import fitz
    try:
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        raise RuntimeError(
            "OCR requires additional packages:\n"
            "  pip install pytesseract Pillow\n"
            "Also install Tesseract from:\n"
            "  https://github.com/UB-Mannheim/tesseract/wiki"
        )

    src = fitz.open(input)
    result = fitz.open()
    inv = 72.0 / dpi

    for page in src:
        pw, ph = page.rect.width, page.rect.height
        mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))

        new_page = result.new_page(width=pw, height=ph)
        new_page.insert_image(new_page.rect, stream=img_bytes)

        data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)

        for i in range(len(data['text'])):
            word = data['text'][i]
            if not word or not word.strip():
                continue
            if int(data['conf'][i]) < 0:
                continue
            x0 = data['left'][i] * inv
            y_top = data['top'][i] * inv
            h_pt = data['height'][i] * inv
            if h_pt < 1:
                continue
            y_baseline = ph - (y_top + h_pt)
            new_page.insert_text(
                fitz.Point(x0, y_baseline),
                word + ' ',
                fontsize=h_pt,
                fontname='helv',
                color=(0, 0, 0),
                render_mode=3,
                overlay=True,
            )

    src.close()
    result.save(output)
    return {"success": True, "pages": len(result)}

def redact_areas(input: str, output: str, page: int, rects: list, **_):
    """Permanently black out the given rectangles on a page (irreversible)."""
    import fitz
    doc = fitz.open(input)
    pg = doc[page]
    for r in rects:
        pg.add_redact_annot(fitz.Rect(r['x0'], r['y0'], r['x1'], r['y1']), fill=(0, 0, 0))
    pg.apply_redactions()
    _save(doc, input, output)
    return {"success": True}

def set_permissions(input: str, output: str,
                    allow_print: bool = True, allow_copy: bool = True,
                    allow_modify: bool = True, allow_annotate: bool = True,
                    allow_forms: bool = True, user_password: str = '', **_):
    """Encrypt with AES-256 and enforce permission flags. Owner password is
    randomly generated so restrictions cannot be removed without the owner key."""
    import fitz, secrets
    perms = fitz.PDF_PERM_ACCESSIBILITY  # always grant accessibility
    if allow_print:    perms |= fitz.PDF_PERM_PRINT | fitz.PDF_PERM_PRINT_HQ
    if allow_copy:     perms |= fitz.PDF_PERM_COPY
    if allow_modify:   perms |= fitz.PDF_PERM_MODIFY | fitz.PDF_PERM_ASSEMBLE
    if allow_annotate: perms |= fitz.PDF_PERM_ANNOTATE
    if allow_forms:    perms |= fitz.PDF_PERM_FORM
    owner_pw = secrets.token_hex(16)
    doc = fitz.open(input)
    _save(doc, input, output,
          encryption=fitz.PDF_ENCRYPT_AES_256,
          owner_pw=owner_pw,
          user_pw=user_password,
          permissions=perms)
    return {"success": True}

def encrypt_pdf(input: str, output: str, password: str, **_):
    import fitz
    doc = fitz.open(input)
    doc.save(output,
             encryption=fitz.PDF_ENCRYPT_AES_256,
             user_pw=password,
             owner_pw=password)
    return {"success": True}

def decrypt_pdf(input: str, output: str, password: str, **_):
    import fitz
    doc = fitz.open(input)
    if doc.needs_pass:
        if not doc.authenticate(password):
            raise ValueError("Incorrect password — could not unlock the PDF.")
    doc.save(output, encryption=fitz.PDF_ENCRYPT_NONE)
    return {"success": True}

def compare_pages(input_a: str, input_b: str, output: str,
                  page_a: int = 0, page_b: int = 0, dpi: int = 150, **_):
    """Render two PDF pages and produce a visual diff image (red highlights = changed pixels)."""
    import fitz
    from PIL import Image, ImageChops
    import io

    mat = fitz.Matrix(dpi / 72, dpi / 72)

    doc_a = fitz.open(input_a)
    pix_a = doc_a[min(page_a, len(doc_a) - 1)].get_pixmap(matrix=mat, alpha=False)
    img_a = Image.open(io.BytesIO(pix_a.tobytes("png"))).convert("RGB")
    doc_a.close()

    doc_b = fitz.open(input_b)
    pix_b = doc_b[min(page_b, len(doc_b) - 1)].get_pixmap(matrix=mat, alpha=False)
    img_b = Image.open(io.BytesIO(pix_b.tobytes("png"))).convert("RGB")
    doc_b.close()

    if img_a.size != img_b.size:
        img_b = img_b.resize(img_a.size, Image.LANCZOS)

    diff = ImageChops.difference(img_a, img_b)
    diff_gray = diff.convert("L")
    mask = diff_gray.point(lambda p: 255 if p > 15 else 0)

    red = Image.new("RGB", img_a.size, (255, 60, 60))
    result = Image.composite(red, img_a, mask)
    result.save(output, format="PNG")

    pixels = list(mask.getdata())
    changed = sum(1 for p in pixels if p > 0)
    total = len(pixels)
    changed_pct = round(changed / total * 100, 1) if total > 0 else 0
    return {"success": True, "changed_pct": changed_pct, "width": img_a.width, "height": img_a.height}

COMMANDS = {
    "to_text": to_text,
    "to_html": to_html,
    "to_image": to_image,
    "to_docx": to_docx,
    "to_xlsx": to_xlsx,
    "merge": merge,
    "split": split,
    "compress": compress,
    "rotate_page": rotate_page,
    "delete_page": delete_page,
    "apply_annotations": apply_annotations,
    "images_to_pdf": images_to_pdf,
    "office_to_pdf": office_to_pdf,
    "url_to_pdf": url_to_pdf,
    "add_watermark": add_watermark,
    "add_stamp": add_stamp,
    "add_header_footer": add_header_footer,
    "add_page_numbers": add_page_numbers,
    "get_outline": get_outline,
    "get_form_fields": get_form_fields,
    "fill_form_fields": fill_form_fields,
    "add_form_field": add_form_field,
    "export_form_data": export_form_data,
    "import_form_data": import_form_data,
    "ocr_pdf": ocr_pdf,
    "redact_areas": redact_areas,
    "set_permissions": set_permissions,
    "encrypt_pdf": encrypt_pdf,
    "decrypt_pdf": decrypt_pdf,
    "compare_pages": compare_pages,
    "reorder_pages": reorder_pages,
    "extract_pages": extract_pages,
    "insert_pages": insert_pages,
    "crop_page": crop_page,
}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No arguments provided"}))
        sys.exit(1)
    try:
        payload = json.loads(sys.argv[1])
        cmd = payload.pop("cmd")
        if cmd not in COMMANDS:
            raise ValueError(f"Unknown command: {cmd}")
        result = COMMANDS[cmd](**payload)
        print(json.dumps(result))
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)
